import os
import re
import json
import hashlib
import logging
import datetime
from typing import List, Dict, Any, Optional, Tuple, Union
from enum import Enum
import pdfplumber
import docx
from io import BytesIO
from app.utils import get_file_extension

logger = logging.getLogger(__name__)

class DocumentType(str, Enum):
    """文件類型枚舉"""
    CONTRACT = "contract"
    LEGAL_OPINION = "legal_opinion"
    LITIGATION = "litigation"
    REGULATION = "regulation"
    MEETING_MINUTES = "meeting_minutes"
    UNKNOWN = "unknown"

def extract_text_from_pdf(file_path: str) -> str:
    """從PDF文件中提取文本

    Args:
        file_path: PDF文件的路徑

    Returns:
        str: 提取的文本內容
    """
    logger.info(f"正在從PDF提取文本: {file_path}")
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n\n"
        return text.strip()
    except Exception as e:
        logger.error(f"PDF文本提取失敗: {str(e)}")
        raise Exception(f"無法處理PDF文件: {str(e)}")

def extract_text_from_docx(file_path: str) -> str:
    """從Word文件中提取文本

    Args:
        file_path: Word文件的路徑

    Returns:
        str: 提取的文本內容
    """
    logger.info(f"正在從DOCX提取文本: {file_path}")
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"DOCX文本提取失敗: {str(e)}")
        raise Exception(f"無法處理DOCX文件: {str(e)}")

def split_into_clauses(document_text: str) -> List[Dict[str, str]]:
    """將文檔文本分割為條款單位

    Args:
        document_text: 文檔的全文本

    Returns:
        List[Dict[str, str]]: 條款列表，每個條款包含ID和內容
    """
    logger.info("正在分割文本為條款")
    
    # 條款標識模式
    # 匹配如"第一條"、"第1條"、"第壹條"等格式，同時處理章節
    clause_pattern = r'(第[一二三四五六七八九十百千萬\d]+[條章節])[：:\s]*(.*?)(?=(第[一二三四五六七八九十百千萬\d]+[條章節])|$)'
    
    # 使用正則表達式找出所有條款
    matches = re.findall(clause_pattern, document_text, re.DOTALL)
    
    if not matches:
        # 如果沒有找到條款，嘗試用段落分割
        paragraphs = [p.strip() for p in document_text.split('\n') if p.strip()]
        return [{"id": f"p{i+1}", "text": p} for i, p in enumerate(paragraphs)]
    
    clauses = []
    for i, match in enumerate(matches):
        clause_id = match[0].strip()
        clause_text = match[1].strip()
        
        # 檢查是否有子條款
        sub_clauses = extract_sub_clauses(clause_text)
        
        if sub_clauses:
            # 如果有子條款，添加主條款和所有子條款
            clauses.append({"id": clause_id, "text": clause_text, "has_sub_clauses": True})
            for sub_id, sub_text in sub_clauses.items():
                clauses.append({
                    "id": f"{clause_id}-{sub_id}",
                    "text": sub_text,
                    "parent_id": clause_id
                })
        else:
            # 如果沒有子條款，只添加主條款
            clauses.append({"id": clause_id, "text": clause_text})
    
    return clauses

def extract_sub_clauses(clause_text: str) -> Dict[str, str]:
    """提取子條款

    Args:
        clause_text: 條款的文本內容

    Returns:
        Dict[str, str]: 子條款ID和內容的字典
    """
    # 匹配如"(一)"、"（一）"等常見的子條款標識
    sub_clause_pattern = r'[（\(]([一二三四五六七八九十]+)[）\)][\s:：]*(.*?)(?=[（\(][一二三四五六七八九十]+[）\)]|$)'
    
    matches = re.findall(sub_clause_pattern, clause_text, re.DOTALL)
    if not matches:
        return {}
    
    sub_clauses = {}
    for match in matches:
        sub_id = match[0].strip()
        sub_text = match[1].strip()
        sub_clauses[sub_id] = sub_text
    
    return sub_clauses

def preprocess_document(file_path: str) -> Dict[str, Any]:
    """文檔預處理主函數

    Args:
        file_path: 文件路徑

    Returns:
        Dict: 包含文檔信息和條款列表的字典
    """
    logger.info(f"開始處理文檔: {file_path}")
    file_extension = get_file_extension(file_path)
    
    # 根據文件類型提取文本
    if file_extension == 'pdf':
        text = extract_text_from_pdf(file_path)
    elif file_extension in ['docx', 'doc']:
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"不支持的文件類型: {file_extension}")
    
    # 分割為條款
    clauses = split_into_clauses(text)
    
    return {
        "document_info": {
            "filename": os.path.basename(file_path),
            "file_type": file_extension,
            "text_length": len(text)
        },
        "full_text": text,
        "clauses": clauses
    }

class DocumentProcessor:
    """文件處理器 - 解析不同類型的法律文件並提取結構化信息"""
    
    def __init__(
        self,
        cache_dir: str = "data/cache",
        templates_dir: str = "data/templates",
        enable_ocr: bool = False
    ):
        """初始化文件處理器
        
        Args:
            cache_dir: 文件處理緩存目錄
            templates_dir: 文件模板目錄
            enable_ocr: 是否啟用OCR功能
        """
        self.cache_dir = cache_dir
        self.templates_dir = templates_dir
        self.enable_ocr = enable_ocr
        
        # 確保目錄存在
        os.makedirs(cache_dir, exist_ok=True)
        
        # 加載文件模板
        self.templates = self._load_templates()
        
        # 正則表達式模式
        self.clause_patterns = {
            "zh-TW": [
                r"第\s*(\d+)\s*條",  # 第N條
                r"(\d+)[\.、]\s*",    # N. 或 N、
                r"（(\d+)）",         # （N）
                r"\((\d+)\)",         # (N)
            ],
            "en": [
                r"Article\s*(\d+)",  # Article N
                r"Section\s*(\d+)",  # Section N
                r"(\d+)[\.]\s*",     # N.
                r"\((\d+)\)",        # (N)
            ]
        }
    
    def _load_templates(self) -> Dict[str, Any]:
        """加載文件模板"""
        templates = {}
        
        if os.path.exists(self.templates_dir):
            for filename in os.listdir(self.templates_dir):
                if filename.endswith(".json"):
                    try:
                        with open(os.path.join(self.templates_dir, filename), "r", encoding="utf-8") as f:
                            template = json.load(f)
                            if "type" in template and "patterns" in template:
                                templates[template["type"]] = template
                    except Exception as e:
                        logger.error(f"加載模板 {filename} 時出錯: {str(e)}")
        
        return templates
    
    async def process_file(
        self,
        file_content: bytes,
        filename: str,
        document_type: Optional[DocumentType] = None
    ) -> Dict[str, Any]:
        """處理文件並提取結構化信息
        
        Args:
            file_content: 文件二進制內容
            filename: 文件名
            document_type: 文件類型(可選)
            
        Returns:
            處理結果，包含文件元數據和提取的結構
        """
        # 計算文件哈希用於緩存
        file_hash = hashlib.md5(file_content).hexdigest()
        cache_path = os.path.join(self.cache_dir, f"{file_hash}.json")
        
        # 檢查緩存
        if os.path.exists(cache_path):
            try:
                with open(cache_path, "r", encoding="utf-8") as f:
                    logger.info(f"從緩存加載文件處理結果: {filename}")
                    return json.load(f)
            except Exception as e:
                logger.warning(f"讀取緩存失敗: {str(e)}")
        
        # 識別文件類型
        file_extension = os.path.splitext(filename)[1].lower()
        
        # 如果沒有指定文件類型，嘗試自動檢測
        if not document_type:
            document_type = await self._detect_document_type(file_content, file_extension)
        
        # 根據文件類型處理
        if file_extension == ".pdf":
            result = await self._process_pdf(file_content, document_type)
        elif file_extension in (".docx", ".doc"):
            result = await self._process_docx(file_content, document_type)
        elif file_extension in (".txt", ".md"):
            result = await self._process_text(file_content, document_type)
        else:
            return {
                "status": "error",
                "message": f"不支持的文件類型: {file_extension}"
            }
        
        # 添加元數據
        result["metadata"] = {
            "filename": filename,
            "file_type": file_extension,
            "document_type": document_type.value if isinstance(document_type, DocumentType) else str(document_type),
            "file_hash": file_hash,
            "processing_timestamp": str(datetime.datetime.now().isoformat())
        }
        
        # 緩存結果
        try:
            with open(cache_path, "w", encoding="utf-8") as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"寫入緩存失敗: {str(e)}")
        
        return result
    
    async def _detect_document_type(
        self, 
        file_content: bytes, 
        file_extension: str
    ) -> DocumentType:
        """檢測文件類型
        
        Args:
            file_content: 文件二進制內容
            file_extension: 文件擴展名
            
        Returns:
            檢測到的文件類型
        """
        # 提取文本用於檢測
        text_sample = ""
        
        try:
            if file_extension == ".pdf":
                with pdfplumber.open(BytesIO(file_content)) as pdf:
                    # 只讀取前幾頁
                    max_pages = min(3, len(pdf.pages))
                    for i in range(max_pages):
                        text_sample += pdf.pages[i].extract_text() or ""
            
            elif file_extension in (".docx", ".doc"):
                doc = docx.Document(BytesIO(file_content))
                # 只讀取前幾段
                max_paras = min(20, len(doc.paragraphs))
                for i in range(max_paras):
                    text_sample += doc.paragraphs[i].text + "\n"
            
            elif file_extension in (".txt", ".md"):
                text_sample = file_content.decode("utf-8", errors="ignore")[:5000]
            
        except Exception as e:
            logger.warning(f"文件類型檢測時出錯: {str(e)}")
            return DocumentType.UNKNOWN
        
        # 使用關鍵詞檢測文件類型
        text_sample = text_sample.lower()
        
        # 檢查模板中的模式
        for doc_type, template in self.templates.items():
            for pattern in template.get("patterns", []):
                if re.search(pattern, text_sample, re.IGNORECASE):
                    return DocumentType(doc_type)
        
        # 基本關鍵詞檢測
        if any(kw in text_sample for kw in ["合約", "合同", "協議", "contract", "agreement"]):
            return DocumentType.CONTRACT
        elif any(kw in text_sample for kw in ["法律意見", "法律分析", "legal opinion"]):
            return DocumentType.LEGAL_OPINION
        elif any(kw in text_sample for kw in ["訴訟", "起訴", "答辯", "litigation", "complaint", "defense"]):
            return DocumentType.LITIGATION
        elif any(kw in text_sample for kw in ["條例", "法規", "辦法", "regulation", "rule"]):
            return DocumentType.REGULATION
        elif any(kw in text_sample for kw in ["會議紀錄", "備忘錄", "minutes", "memorandum"]):
            return DocumentType.MEETING_MINUTES
        
        # 默認類型
        return DocumentType.UNKNOWN
    
    async def _process_pdf(
        self, 
        file_content: bytes,
        document_type: DocumentType
    ) -> Dict[str, Any]:
        """處理PDF文件
        
        Args:
            file_content: PDF文件二進制內容
            document_type: 文件類型
            
        Returns:
            處理結果
        """
        result = {
            "status": "success",
            "document_type": document_type.value if isinstance(document_type, DocumentType) else str(document_type),
            "pages": [],
            "structure": {}
        }
        
        try:
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                # 處理每一頁
                for i, page in enumerate(pdf.pages):
                    # 提取文本
                    text = page.extract_text() or ""
                    
                    # 提取表格
                    tables = []
                    for table in page.extract_tables():
                        if table:
                            tables.append({
                                "headers": table[0] if table and len(table) > 0 else [],
                                "rows": table[1:] if table and len(table) > 1 else []
                            })
                    
                    result["pages"].append({
                        "page_number": i + 1,
                        "text": text,
                        "tables": tables
                    })
                
                # 根據文件類型處理結構
                if document_type == DocumentType.CONTRACT:
                    # 使用原有函數處理合約內容
                    all_text = "\n".join([page.get("text", "") for page in result["pages"]])
                    clauses = split_into_clauses(all_text)
                    result["structure"] = {
                        "document_info": {
                            "file_type": "pdf",
                            "text_length": len(all_text)
                        },
                        "full_text": all_text,
                        "clauses": clauses
                    }
                else:
                    # 通用結構提取
                    all_text = "\n".join([page.get("text", "") for page in result["pages"]])
                    result["structure"] = {
                        "full_text": all_text,
                        "paragraphs": [p.strip() for p in all_text.split('\n') if p.strip()]
                    }
        
        except Exception as e:
            logger.error(f"處理PDF時出錯: {str(e)}")
            return {
                "status": "error",
                "message": f"PDF處理錯誤: {str(e)}"
            }
        
        return result
    
    async def _process_docx(
        self, 
        file_content: bytes,
        document_type: DocumentType
    ) -> Dict[str, Any]:
        """處理DOCX文件
        
        Args:
            file_content: DOCX文件二進制內容
            document_type: 文件類型
            
        Returns:
            處理結果
        """
        result = {
            "status": "success",
            "document_type": document_type.value if isinstance(document_type, DocumentType) else str(document_type),
            "paragraphs": [],
            "structure": {}
        }
        
        try:
            doc = docx.Document(BytesIO(file_content))
            
            # 提取段落
            full_text = []
            for para in doc.paragraphs:
                result["paragraphs"].append({
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal"
                })
                full_text.append(para.text)
            
            # 提取表格
            result["tables"] = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                result["tables"].append({
                    "headers": table_data[0] if table_data else [],
                    "rows": table_data[1:] if len(table_data) > 1 else []
                })
            
            # 根據文件類型處理結構
            if document_type == DocumentType.CONTRACT:
                # 使用原有函數處理合約內容
                all_text = "\n".join(full_text)
                clauses = split_into_clauses(all_text)
                result["structure"] = {
                    "document_info": {
                        "file_type": "docx",
                        "text_length": len(all_text)
                    },
                    "full_text": all_text,
                    "clauses": clauses
                }
            else:
                # 通用結構提取
                all_text = "\n".join(full_text)
                result["structure"] = {
                    "full_text": all_text,
                    "paragraphs": full_text
                }
        
        except Exception as e:
            logger.error(f"處理DOCX時出錯: {str(e)}")
            return {
                "status": "error",
                "message": f"DOCX處理錯誤: {str(e)}"
            }
        
        return result
    
    async def _process_text(
        self, 
        file_content: bytes,
        document_type: DocumentType
    ) -> Dict[str, Any]:
        """處理純文本文件
        
        Args:
            file_content: 文本文件二進制內容
            document_type: 文件類型
            
        Returns:
            處理結果
        """
        result = {
            "status": "success",
            "document_type": document_type.value if isinstance(document_type, DocumentType) else str(document_type),
            "content": file_content.decode("utf-8", errors="ignore"),
            "structure": {}
        }
        
        try:
            text = result["content"]
            lines = text.split("\n")
            
            # 根據文件類型處理結構
            if document_type == DocumentType.CONTRACT:
                # 使用原有函數處理合約內容
                clauses = split_into_clauses(text)
                result["structure"] = {
                    "document_info": {
                        "file_type": "txt",
                        "text_length": len(text)
                    },
                    "full_text": text,
                    "clauses": clauses
                }
            else:
                # 通用結構提取
                result["structure"] = {
                    "full_text": text,
                    "paragraphs": [line for line in lines if line.strip()]
                }
        
        except Exception as e:
            logger.error(f"處理文本文件時出錯: {str(e)}")
            return {
                "status": "error",
                "message": f"文本處理錯誤: {str(e)}"
            }
        
        return result
    
    def extract_contract_parties(self, text: str) -> List[Dict[str, str]]:
        """提取合約當事人
        
        Args:
            text: 合約文本
            
        Returns:
            List[Dict[str, str]]: 合約當事人列表
        """
        parties = []
        
        # 尋找當事人部分
        party_patterns = [
            r"甲\s*方[：:]\s*([^\n,，、]+)",
            r"乙\s*方[：:]\s*([^\n,，、]+)",
            r"丙\s*方[：:]\s*([^\n,，、]+)",
            r"買\s*方[：:]\s*([^\n,，、]+)",
            r"賣\s*方[：:]\s*([^\n,，、]+)",
            r"出租\s*人[：:]\s*([^\n,，、]+)",
            r"承租\s*人[：:]\s*([^\n,，、]+)",
            r"Party\s*A[：:]\s*([^\n,，、]+)",
            r"Party\s*B[：:]\s*([^\n,，、]+)"
        ]
        
        for pattern in party_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                party_name = match.group(1).strip()
                party_type = match.group(0).split("[：:]")[0].strip()
                parties.append({
                    "type": party_type,
                    "name": party_name
                })
        
        return parties
    
    def extract_contract_date(self, text: str) -> str:
        """提取合約日期
        
        Args:
            text: 合約文本
            
        Returns:
            str: 合約日期
        """
        # 嘗試多種日期格式
        date_patterns = [
            r"中華民國\s*(\d+)\s*年\s*(\d+)\s*月\s*(\d+)\s*日",
            r"西元\s*(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日",
            r"(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日",
            r"(\d{4})[./](\d{1,2})[./](\d{1,2})",
            r"(\d{1,2})[./](\d{1,2})[./](\d{4})"
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, text)
            if match:
                # 根據格式返回標準日期
                if "中華民國" in pattern:
                    year = int(match.group(1)) + 1911  # 民國年轉換為西元年
                    return f"{year}/{match.group(2)}/{match.group(3)}"
                else:
                    return f"{match.group(1)}/{match.group(2)}/{match.group(3)}"
        
        return ""
